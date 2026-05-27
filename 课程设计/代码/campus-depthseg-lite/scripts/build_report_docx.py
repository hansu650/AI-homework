"""Build the course-design report draft in Markdown and DOCX formats."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image

PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = PROJECT_ROOT / "REPORT_DRAFT.md"
REPORT_DOCX = PROJECT_ROOT / "课程设计报告_初稿.docx"


@dataclass(frozen=True)
class Figure:
    number: str
    title: str
    path: str
    width_inches: float = 5.8
    max_height_inches: float = 6.7
    page_break_before: bool = False

    @property
    def exists(self) -> bool:
        return (PROJECT_ROOT / self.path).exists()


EXPERIMENT_ROWS = [
    ["RGB-only", "RGB", "none", "0.4226", "0.5964", "0.5881", "1.5588"],
    ["RGBD-concat", "RGB + Depth", "input concat", "0.4683", "0.6318", "0.6294", "1.4205"],
    ["RGBD-boundary", "RGB + Depth", "depth boundary residual fusion", "0.4206", "0.5974", "0.5925", "1.5317"],
    [
        "RGBD-concat-boundary",
        "RGB + Depth",
        "input concat + boundary residual fusion",
        "0.4605",
        "0.6271",
        "0.6191",
        "1.4378",
    ],
]

PER_CLASS_ROWS = [
    ["RGB-only", "0.3549", "0.4731", "0.5495", "0.4239", "0.3114"],
    ["RGBD-concat", "0.3857", "0.5751", "0.5647", "0.4722", "0.3437"],
    ["RGBD-boundary", "0.3476", "0.4629", "0.5443", "0.4276", "0.3205"],
    ["RGBD-concat-boundary", "0.3636", "0.5530", "0.5595", "0.4899", "0.3366"],
]

COMPLEXITY_ROWS = [
    ["rgb", "1,511,689", "1.5117"],
    ["rgbd_concat", "1,512,121", "1.5121"],
    ["rgbd_boundary", "1,512,413", "1.5124"],
    ["rgbd_concat_boundary", "1,512,845", "1.5128"],
]

ASSET_FIGURES = [
    Figure("图3-1", "NYU5 五类标签像素分布", "outputs/figures/nyu5_class_distribution.png", 5.6),
    Figure("图3-2", "训练集样例可视化", "outputs/figures/nyu5_gallery_train.png", 5.7),
    Figure("图3-3", "验证集样例可视化", "outputs/figures/nyu5_gallery_val.png", 5.7),
    Figure("图3-4", "测试集样例可视化", "outputs/figures/nyu5_gallery_test.png", 5.7),
    Figure("图5-1", "RGBD-concat 训练损失曲线", "outputs/runs/exp02_rgbd_concat_e20/training_loss_curve.png", 5.4),
    Figure("图5-2", "RGBD-concat 验证 mIoU 曲线", "outputs/runs/exp02_rgbd_concat_e20/validation_miou_curve.png", 5.4),
    Figure("图5-3", "RGBD-concat 验证 Pixel Acc 曲线", "outputs/runs/exp02_rgbd_concat_e20/validation_pixel_acc_curve.png", 5.4),
    Figure(
        "图5-4",
        "四种方法在同一批测试样本上的预测对比",
        "outputs/report_assets/method_comparison/method_comparison_grid.png",
        6.2,
        page_break_before=True,
    ),
    Figure("图5-5", "最佳模型归一化混淆矩阵", "outputs/report_assets/confusion_matrix/confusion_matrix_normalized.png", 5.3),
    Figure("图5-6", "测试集中表现较好的样例", "outputs/report_assets/error_cases/best_cases.png", 6.0),
    Figure("图5-7", "测试集中表现较差的样例", "outputs/report_assets/error_cases/worst_cases.png", 6.0),
    Figure("图5-8", "自采集校园场景定性展示", "outputs/report_assets/campus_demo/campus_demo_gallery.png", 6.3),
]

REFERENCES = [
    "[1] Silberman N, Hoiem D, Kohli P, et al. Indoor segmentation and support inference from RGBD images[C]//European Conference on Computer Vision. Berlin: Springer, 2012: 746-760.",
    "[2] Long J, Shelhamer E, Darrell T. Fully convolutional networks for semantic segmentation[C]//IEEE Conference on Computer Vision and Pattern Recognition. 2015: 3431-3440.",
    "[3] Ronneberger O, Fischer P, Brox T. U-Net: convolutional networks for biomedical image segmentation[C]//International Conference on Medical Image Computing and Computer-Assisted Intervention. 2015: 234-241.",
    "[4] Chen L C, Zhu Y, Papandreou G, et al. Encoder-decoder with atrous separable convolution for semantic image segmentation[C]//European Conference on Computer Vision. 2018: 801-818.",
    "[5] Lin T Y, Dollar P, Girshick R, et al. Feature pyramid networks for object detection[C]//IEEE Conference on Computer Vision and Pattern Recognition. 2017: 2117-2125.",
    "[6] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]//IEEE Conference on Computer Vision and Pattern Recognition. 2016: 770-778.",
    "[7] Badrinarayanan V, Kendall A, Cipolla R. SegNet: a deep convolutional encoder-decoder architecture for image segmentation[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2017, 39(12): 2481-2495.",
    "[8] Hazirbas C, Ma L, Domokos C, et al. FuseNet: incorporating depth into semantic segmentation via fusion-based CNN architecture[C]//Asian Conference on Computer Vision. 2016: 213-228.",
    "[9] Seichter D, Kohler M, Lewandowski B, et al. Efficient RGB-D semantic segmentation for indoor scene analysis[C]//IEEE International Conference on Robotics and Automation. 2021: 13525-13531.",
    "[10] Paszke A, Gross S, Massa F, et al. PyTorch: an imperative style, high-performance deep learning library[C]//Advances in Neural Information Processing Systems. 2019: 8024-8035.",
    "[11] Kingma D P, Ba J. Adam: a method for stochastic optimization[C]//International Conference on Learning Representations. 2015.",
    "[12] Milletari F, Navab N, Ahmadi S A. V-Net: fully convolutional neural networks for volumetric medical image segmentation[C]//International Conference on 3D Vision. 2016: 565-571.",
]


def main() -> None:
    markdown = build_markdown()
    REPORT_MD.write_text(markdown, encoding="utf-8")
    build_docx(markdown)
    print(f"markdown: {REPORT_MD}")
    print(f"docx: {REPORT_DOCX}")


def build_markdown() -> str:
    lines: list[str] = []
    lines.extend(
        [
            "# 面向校园室内空间巡检的轻量 RGB-D 语义分割与空间占用分析系统",
            "",
            "课程名称：《人工智能技术与应用》课程设计",
            "",
            "课设题目：面向校园室内空间巡检的轻量 RGB-D 语义分割与空间占用分析系统",
            "",
            "姓名：易唯",
            "",
            "学号：202231116020106",
            "",
            "专业年级：计算机科学与技术2022级",
            "",
            "指导教师/职称：王雷春/副教授",
            "",
            "日期：2026年6月",
            "",
            "## 摘 要",
            "",
            "针对校园教室、实验室、走廊等室内空间中通道遮挡、物品堆放和人工巡检主观性较强的问题，本文设计并实现了一套面向校园室内空间巡检的轻量 RGB-D 语义分割与空间占用分析系统。系统以 NYUDepthV2 数据集为基础，将原始室内场景标签映射为 other、floor、wall、obstacle、door_window 五类校园空间语义，并构建 RGB-only、RGBD-concat、RGBD-boundary 和 RGBD-concat-boundary 四种模型变体进行对比。模型采用轻量四阶段编码器、Weighted-FPN 多尺度解码器以及可选的深度边界残差融合模块，训练损失由交叉熵损失和 Dice 损失组成。",
            "",
            "实验结果表明，在 218 张测试图像上，RGBD-concat 取得最佳整体表现，mIoU 为 0.4683，Pixel Acc 为 0.6318，Mean Acc 为 0.6294；相较 RGB-only，mIoU 提升 4.57 个百分点，说明完整 depth 信息对室内空间分割具有积极作用。RGBD-concat-boundary 的 obstacle IoU 最高，为 0.4899，表明深度边界先验对障碍区域具有一定帮助，但其整体 mIoU 未超过 RGBD-concat，说明简单 Sobel 深度边界残差注入仍存在噪声、冗余或融合方式不足。系统进一步实现了基于预测 mask 的空间占用比例、障碍区域框和风险提示可视化，并在自采集校园 RGB 图像上进行了定性演示。自采集图像没有像素级标签和真实 depth，仅用于展示系统流程，不参与训练和定量评价。",
            "",
            "关键词：RGB-D 语义分割；校园巡检；空间占用分析；轻量网络；多尺度融合",
            "",
            "## 目录",
            "",
            "[Word 中更新自动目录或手动整理目录]",
            "",
        ]
    )
    add_section_one(lines)
    add_section_two(lines)
    add_section_three(lines)
    add_section_four(lines)
    add_section_five(lines)
    add_section_six(lines)
    add_references(lines)
    add_appendices(lines)
    return "\n".join(lines) + "\n"


def add_section_one(lines: list[str]) -> None:
    lines.extend(
        [
            "## 1 引言",
            "",
            "### 1.1 背景与意义",
            "",
            "校园室内空间包含教室、实验室、宿舍楼公共区域、图书馆和走廊等多类场景。这些区域在日常使用中可能出现临时堆放物品、通道遮挡、门窗区域被遮挡、地面可见区域不足等现象。传统人工巡检通常依赖巡检人员的主观观察，记录粒度不稳定，难以长期形成结构化的空间占用分析结果。随着计算机视觉和深度学习的发展，利用图像语义分割方法自动识别室内空间中的地面、墙面、障碍物和门窗区域，可以为校园室内空间巡检提供更直观的辅助信息。",
            "",
            "与普通图像分类不同，语义分割需要对每个像素进行类别预测，因此更适合描述空间区域的占用关系。RGB 图像包含纹理和颜色信息，depth 图像包含几何距离信息，二者结合能够更好地区分地面、墙面和家具等室内结构。对于校园巡检场景而言，地面可见率和下半区域障碍物比例可以进一步转化为空间占用分析指标，为“通道是否可能被遮挡”提供启发式判断。本文围绕这一应用目标，构建轻量 RGB-D 语义分割模型，并在其输出基础上进行空间占用分析和风险提示。",
            "",
            "### 1.2 国内外研究现状",
            "",
            "语义分割任务从早期基于手工特征和传统分类器的方法，逐渐发展到以全卷积网络、编码器-解码器结构和多尺度特征融合为主的深度学习方法。FCN 首次将分类网络改造为端到端像素预测网络，U-Net 通过跳跃连接增强了编码器和解码器之间的细节传递，DeepLab 系列利用空洞卷积和解码结构改善多尺度上下文建模。FPN 则通过自顶向下的多尺度融合结构提升了不同尺度目标的表达能力。",
            "",
            "在 RGB-D 室内语义分割方向，NYUDepthV2 是常用公开数据集之一，提供 RGB 图像、深度信息和语义标签。相关研究表明，depth 可以补充 RGB 难以表达的空间几何结构，有助于区分地面、墙面、家具和门窗等区域。然而，深度信息的使用方式并不唯一：一种简单方式是将 RGB 与 depth 直接拼接输入网络，另一类方法则尝试提取深度边界、几何先验或跨模态注意力进行融合。对于课程设计而言，在保持模型简洁、可读和可复现的前提下比较多种 RGB-D 融合方式，比盲目引入复杂大模型更符合任务目标。",
            "",
            "### 1.3 本文主要工作",
            "",
            "本文主要完成以下工作：第一，基于 NYUDepthV2 数据集构建五类校园室内空间标签体系，将原始室内物体和结构类别映射为 other、floor、wall、obstacle、door_window。第二，设计轻量语义分割模型 CampusDepthSegLite，模型由轻量四阶段编码器、Weighted-FPN 解码器和可选深度边界残差融合模块构成。第三，设置 RGB-only、RGBD-concat、RGBD-boundary 和 RGBD-concat-boundary 四种实验变体，比较 RGB 信息、完整 depth 信息和深度边界信息的作用。第四，在预测 mask 基础上实现空间占用分析，包括地面可见率、障碍物占比、下半区域占用比例、风险分数和障碍区域框可视化。第五，在自采集校园 RGB 场景上进行定性演示，展示系统从图像输入到语义预测、占用图、风险框和文字摘要的完整流程。",
            "",
            "### 1.4 组织结构",
            "",
            "本文共分为六章。第一章介绍研究背景、意义和主要工作；第二章介绍语义分割、RGB-D 多模态信息和多尺度解码等相关理论；第三章给出问题定义、数据集和预处理方法；第四章介绍模型结构和系统实现；第五章展示实验结果、消融分析、可视化和自采集校园场景定性演示；第六章总结工作并讨论局限性与未来改进方向。",
            "",
        ]
    )


def add_section_two(lines: list[str]) -> None:
    lines.extend(
        [
            "## 2 相关理论与技术",
            "",
            "### 2.1 语义分割任务",
            "",
            "语义分割是计算机视觉中的像素级分类任务。给定输入图像，模型需要为每个像素预测一个语义类别，从而得到与原图空间尺寸对应的类别 mask。与图像分类只输出整幅图像类别不同，语义分割能够保留目标区域的空间形状和边界信息，适合用于室内空间区域分析。本文关注的输出类别包括地面、墙面、障碍物、门窗区域以及其他区域，预测结果不仅用于评价 mIoU、Pixel Acc 等指标，也用于后续空间占用比例计算。",
            "",
            "### 2.2 RGB-D 多模态信息",
            "",
            "RGB 图像提供颜色、纹理和外观信息，但在光照变化、遮挡和材质相似的室内场景中，单纯依赖 RGB 容易产生混淆。Depth 图像提供相对几何信息，能够反映物体与相机之间的空间关系，对区分地面、墙面、家具等具有帮助。本文在实验中比较了 RGB-only 和 RGBD-concat，验证完整 depth 直接输入对整体分割性能的提升；同时比较 RGBD-boundary 和 RGBD-concat-boundary，用于分析深度边界先验是否能够进一步改善障碍区域识别。",
            "",
            "### 2.3 轻量多尺度特征提取",
            "",
            "室内场景中既有墙面、地面等大面积结构，也有椅子、柜子、门窗等局部区域，因此模型需要同时表达局部细节和全局布局。本文使用四阶段轻量编码器，默认通道为 [48, 96, 192, 384]。每个阶段通过下采样卷积逐步降低空间分辨率，并通过轻量残差块提取特征。该设计没有引入复杂 attention 或大规模预训练骨干，参数量约 1.51M，便于课程设计环境下训练、测试和解释。",
            "",
            "### 2.4 FPN 多尺度解码",
            "",
            "Feature Pyramid Network 的核心思想是融合不同尺度的特征，使高层语义信息和低层空间细节共同参与预测。本文的 Weighted-FPN 解码器先将四个尺度特征统一到 128 通道，再上采样到最高分辨率特征尺度，使用可学习权重进行归一化加权融合，最后通过卷积平滑并输出五类分割 logits。该结构相比简单单尺度解码更适合室内场景中不同大小区域的分割。",
            "",
            "### 2.5 损失函数与评价指标",
            "",
            "训练损失采用 CrossEntropyLoss 与 DiceLoss 的组合。交叉熵损失适合像素级多分类，DiceLoss 能从区域重叠角度缓解类别面积差异带来的影响。评价指标包括 Pixel Acc、Mean Acc、per-class IoU 和 mean IoU。Pixel Acc 衡量整体像素预测正确率，Mean Acc 对各类别准确率求平均，IoU 衡量预测区域和真实区域的交并比，mIoU 是语义分割任务中常用的综合指标。本文在测试集上以 mIoU 作为主要模型比较指标，同时结合 per-class IoU 分析不同类别的表现。",
            "",
        ]
    )


def add_section_three(lines: list[str]) -> None:
    lines.extend(
        [
            "## 3 问题定义与数据集",
            "",
            "### 3.1 问题定义",
            "",
            "本文任务定义为校园室内空间的轻量 RGB-D 语义分割与占用分析。对于 NYUDepthV2 实验，输入为 RGB 图像和对应 depth 图像，输出为尺寸相同的五类语义分割 mask。对于自采集校园场景演示，由于没有真实 depth 和像素级标签，系统仅使用 RGB-only checkpoint 生成定性预测，不参与训练和定量评价。分割 mask 进一步被用于计算地面比例、障碍物比例、下半区域地面比例和风险分数，从而形成可视化占用分析结果。",
            "",
            "### 3.2 NYUDepthV2 数据集",
            "",
            "NYUDepthV2 是室内 RGB-D 语义理解常用公开数据集，包含多种室内场景的 RGB 图像、深度图和像素级语义标签。本文从 nyu_depth_v2_labeled.mat 导出 1449 张 RGB-D 样本，统一调整输入尺寸为 240 x 320，并划分为训练集 1014 张、验证集 217 张、测试集 218 张。数据集用于训练和定量评价四种模型变体。",
            "",
            "### 3.3 五类校园空间标签映射",
            "",
            "原始 NYUDepthV2 标签类别较多，直接用于课程设计会使任务过于分散。结合校园室内巡检需求，本文将标签映射为五类：0 other、1 floor、2 wall、3 obstacle、4 door_window。其中 floor 表示地面区域，wall 表示墙面和天花板等大面积结构，obstacle 表示可能影响通行或空间占用的桌椅、柜子、床、箱包等物体，door_window 表示门窗及相关遮挡结构，other 表示不属于上述类别的区域。",
            "",
            "### 3.4 数据预处理与划分",
            "",
            "数据预处理包括 RGB 图像读取、depth 读取与归一化、标签映射、尺寸调整、随机水平翻转和 RGB 标准化。训练、验证和测试 split 使用相对路径记录，便于项目在不同机器上复现。Depth 保存为 16-bit PNG 以保留深度范围，标签保存为 8-bit PNG，取值为 0 到 4 或 ignore index。所有几何变换在 RGB、depth 和 label 上同步执行，保证像素对应关系不被破坏。",
            "",
            "### 3.5 类别分布分析",
            "",
            "五类标签分布反映了室内场景的明显不均衡性。墙面、地面和障碍物通常占据较大区域，而 door_window 的面积相对较小，容易受到遮挡、光照和边界模糊影响。类别不均衡也是 door_window IoU 较低的重要原因之一。本文在报告中使用类别分布图和 train/val/test 样例图展示数据预处理结果。",
            "",
            figure_md("图3-1", "NYU5 五类标签像素分布", "outputs/figures/nyu5_class_distribution.png"),
            "",
            figure_md("图3-2", "训练集样例可视化", "outputs/figures/nyu5_gallery_train.png"),
            "",
            figure_md("图3-3", "验证集样例可视化", "outputs/figures/nyu5_gallery_val.png"),
            "",
            figure_md("图3-4", "测试集样例可视化", "outputs/figures/nyu5_gallery_test.png"),
            "",
        ]
    )


def add_section_four(lines: list[str]) -> None:
    lines.extend(
        [
            "## 4 模型设计与系统实现",
            "",
            "### 4.1 系统总体流程",
            "",
            "系统流程包括数据准备、模型推理、语义分割、空间占用分析和可视化输出五个阶段。训练和定量评价阶段使用 NYUDepthV2 导出的 RGB-D 数据；自采集校园演示阶段仅使用 RGB 图像和 RGB-only checkpoint。模型输出五类分割 mask 后，系统计算 floor_ratio、lower_floor_ratio、obstacle_ratio、lower_obstacle_ratio 和 risk_score，并利用连通域分析从 obstacle mask 中得到风险区域框。最终输出包括 Prediction mask、Occupancy map、Risk boxes、Metrics block 和 Text summary。",
            "",
            "[待插入：系统总体流程图]",
            "",
            "### 4.2 CampusDepthSegLite 模型结构",
            "",
            "CampusDepthSegLite 是本文实现的轻量语义分割模型。模型编码器为四阶段层次结构，默认通道数为 [48, 96, 192, 384]，每个阶段包含下采样卷积和轻量残差块。残差块由 depthwise convolution、pointwise convolution、normalization 和 ReLU 组成，用于在较低计算量下提取局部特征。解码器使用 Weighted-FPN，将四个尺度特征统一到 128 通道后进行上采样、加权融合和平滑卷积，最后输出五类 segmentation logits。",
            "",
            "### 4.3 RGB-D 融合策略",
            "",
            "为比较不同深度信息使用方式，本文设置四个实验变体。RGB-only 只使用 RGB 图像，是视觉基线；RGBD-concat 将 RGB 与 depth 直接拼接为四通道输入，用于验证完整 depth 信息是否有效；RGBD-boundary 使用 RGB 编码器，并将 depth 经过 Sobel 边界提取后残差注入多尺度特征；RGBD-concat-boundary 同时使用完整 depth 拼接和 depth boundary residual fusion，用于验证二者是否互补。",
            "",
            "### 4.4 深度边界残差融合模块",
            "",
            "深度边界残差融合模块使用固定 Sobel 卷积从 depth 中提取边界响应，并将边界图 resize 到每个编码器阶段对应的空间尺度。随后使用 1 x 1 卷积将边界响应映射到对应通道数，并通过可学习标量 alpha 残差加入特征。该模块的设计目标是将深度几何边界作为辅助提示，增强障碍物与结构边界的表达。不过实验结果显示，简单边界残差注入并不一定提升整体 mIoU，说明边界噪声和融合方式仍需进一步改进。",
            "",
            "### 4.5 空间占用分析模块",
            "",
            "空间占用分析模块直接基于预测 mask 工作，不使用目标检测模型。系统将 floor 类视为可通行区域的重要线索，将 obstacle 类视为可能影响通行的区域。风险分数定义为 risk_score = 0.45 x (1 - lower_floor_ratio) + 0.35 x lower_obstacle_ratio + 0.20 x obstacle_ratio。该分数是启发式指标，只用于课程设计演示和可视化辅助，不能直接替代真实安全巡检决策。",
            "",
            "### 4.6 实验环境与超参数",
            "",
            "实验使用 PyTorch 和 Lightning 实现。训练优化器为 AdamW，学习率为 1e-4，batch size 为 4，每个实验训练 20 epoch。模型 checkpoint 监控指标为 val_mIoU，mode 设置为 max。损失函数为 CrossEntropyLoss + DiceLoss。输入尺寸统一为 240 x 320。模型复杂度约为 1.51M 参数，适合课程设计环境下进行完整训练、评估和可视化分析。",
            "",
            table_md("表4-1", "模型复杂度统计", ["Variant", "Params", "Params (M)"], COMPLEXITY_ROWS),
            "",
        ]
    )


def add_section_five(lines: list[str]) -> None:
    lines.extend(
        [
            "## 5 实验结果与分析",
            "",
            "### 5.1 训练过程分析",
            "",
            "训练过程以 RGBD-concat 作为定量实验中的最佳整体模型进行重点观察。从训练损失曲线可以看到，模型在训练过程中逐步收敛；验证 mIoU 和 Pixel Acc 曲线用于观察模型在验证集上的泛化趋势。由于课程设计模型较轻量，训练曲线可能存在一定波动，但整体结果能够支持不同融合策略之间的比较。",
            "",
            figure_md("图5-1", "RGBD-concat 训练损失曲线", "outputs/runs/exp02_rgbd_concat_e20/training_loss_curve.png"),
            "",
            figure_md("图5-2", "RGBD-concat 验证 mIoU 曲线", "outputs/runs/exp02_rgbd_concat_e20/validation_miou_curve.png"),
            "",
            figure_md("图5-3", "RGBD-concat 验证 Pixel Acc 曲线", "outputs/runs/exp02_rgbd_concat_e20/validation_pixel_acc_curve.png"),
            "",
            "### 5.2 主实验结果",
            "",
            "四个实验结果如表5-1所示。RGBD-concat 在测试集上取得最高 mIoU，为 0.4683，Pixel Acc 为 0.6318，Mean Acc 为 0.6294，Test Loss 为 1.4205。相比 RGB-only，RGBD-concat 的 mIoU 提升 4.57 个百分点，说明完整 depth 信息能够有效补充 RGB 图像中的几何线索。RGBD-boundary 的整体 mIoU 为 0.4206，略低于 RGB-only；RGBD-concat-boundary 的 mIoU 为 0.4605，接近但未超过 RGBD-concat。",
            "",
            table_md(
                "表5-1",
                "四种模型变体测试集结果",
                ["Method", "Input", "Fusion", "mIoU", "Pixel Acc", "Mean Acc", "Test Loss"],
                EXPERIMENT_ROWS,
            ),
            "",
            "### 5.3 分类别 IoU 分析",
            "",
            "从 per-class IoU 可以看出，RGBD-concat 在 other、floor、wall、door_window 上均取得较好表现，其中 floor IoU 达到 0.5751，说明 depth 对地面区域具有明显帮助。RGBD-concat-boundary 的 obstacle IoU 最高，为 0.4899，说明深度边界信息对障碍区域具有一定补充作用。然而 door_window 类整体 IoU 较低，最高仅为 0.3437。这可能与门窗区域面积较小、外观变化大、与墙面边界接近以及光照影响有关。",
            "",
            table_md(
                "表5-2",
                "四种模型变体 per-class IoU",
                ["Method", "other", "floor", "wall", "obstacle", "door_window"],
                PER_CLASS_ROWS,
            ),
            "",
            "### 5.4 消融实验讨论",
            "",
            "消融实验说明，完整 depth 信息比单纯深度边界残差信息更有助于整体分割性能。RGBD-concat 的整体 mIoU 最高，说明直接输入 depth 能让编码器学习更完整的几何分布。RGBD-boundary 未超过 RGB-only，说明只使用 Sobel 边界可能丢失深度区域信息，并可能引入噪声。RGBD-concat-boundary 在 obstacle IoU 上表现最好，但整体 mIoU 略低于 RGBD-concat，说明边界信息对障碍物有帮助，但当前残差注入方式还没有稳定提升所有类别。",
            "",
            "### 5.5 可视化结果分析",
            "",
            "方法对比图展示了同一批测试样本在四种模型下的预测结果。整体上，RGBD-concat 对地面和墙面的连续区域预测更稳定；RGBD-concat-boundary 在部分障碍区域上给出更明显的响应，但也可能带来局部边界噪声。可视化结果与定量指标一致：depth 信息有效，边界信息对 obstacle 有帮助但不是整体最优。",
            "",
            figure_md("图5-4", "四种方法在同一批测试样本上的预测对比", "outputs/report_assets/method_comparison/method_comparison_grid.png"),
            "",
            "### 5.6 混淆矩阵与错误分析",
            "",
            "混淆矩阵用于观察类别之间的混淆情况。door_window 类容易与 wall 和 other 混淆，主要原因是门窗区域常与墙面相邻，边界较细，且受光照影响明显。错误案例显示，在遮挡严重、物体堆叠或光照不均的场景中，模型容易将部分障碍物预测为 other 或 wall。较好案例通常具有清晰的地面区域、较规整的室内布局和较明显的障碍物边界。",
            "",
            figure_md("图5-5", "最佳模型归一化混淆矩阵", "outputs/report_assets/confusion_matrix/confusion_matrix_normalized.png"),
            "",
            figure_md("图5-6", "测试集中表现较好的样例", "outputs/report_assets/error_cases/best_cases.png"),
            "",
            figure_md("图5-7", "测试集中表现较差的样例", "outputs/report_assets/error_cases/worst_cases.png"),
            "",
            "### 5.7 自采集校园场景定性展示",
            "",
            "为展示系统在真实校园室内图像上的可视化流程，本文使用图书馆、实验室、教室、走廊和遮挡较多室内场景中的自采集 RGB 图像进行定性演示。需要强调的是，自采集图片没有像素级 GT，也没有真实 depth，因此该部分不参与训练、验证和测试，不计算 mIoU、Pixel Acc 或 Mean Acc。该部分使用 RGB-only checkpoint 进行推理，仅展示 Prediction mask、Occupancy map、Risk boxes、Metrics block 和 Text summary 等系统输出流程。自采集校园图像仅用于真实场景下的定性展示，不作为模型泛化性能的定量证明。",
            "",
            figure_md("图5-8", "自采集校园场景定性展示", "outputs/report_assets/campus_demo/campus_demo_gallery.png"),
            "",
            "从图5-8可以看出，系统能够根据 RGB-only 模型输出生成占用图、风险框和文字摘要。由于自采集图像与 NYUDepthV2 存在 domain gap，且缺少真实 depth，预测结果并不等同于 RGB-D 定量实验结果。该演示的价值在于展示系统流程和交互形态，而不是证明真实校园场景下的定量精度。",
            "",
        ]
    )


def add_section_six(lines: list[str]) -> None:
    lines.extend(
        [
            "## 6 总结与展望",
            "",
            "### 6.1 工作总结",
            "",
            "本文围绕校园室内空间巡检需求，完成了轻量 RGB-D 语义分割与空间占用分析系统的设计与实现。系统将 NYUDepthV2 映射为五类校园空间标签，构建了 CampusDepthSegLite 模型，并比较了 RGB-only、RGBD-concat、RGBD-boundary 和 RGBD-concat-boundary 四种融合策略。实验表明，RGBD-concat 在整体 mIoU 上表现最好，说明完整 depth 信息对室内空间分割有效；RGBD-concat-boundary 在 obstacle IoU 上表现最好，说明深度边界对障碍区域有一定帮助。系统还实现了基于分割 mask 的空间占用分析、障碍区域框和风险提示可视化，并完成自采集校园 RGB 图像的定性演示。",
            "",
            "### 6.2 不足与局限性",
            "",
            "本文仍存在若干局限。第一，NYUDepthV2 与真实校园场景之间存在 domain gap，模型在自采集校园图像上的效果不能直接等同于测试集指标。第二，自采集图像没有像素级标签和真实 depth，只能进行定性展示，不能作为定量评价。第三，深度边界残差融合使用 Sobel 边界和简单残差注入，可能受到深度噪声影响，整体 mIoU 未超过 RGBD-concat。第四，风险评分是启发式规则，不能直接用于真实安全巡检决策。第五，door_window 类样本面积较小且易与 wall、other 混淆，仍需更有针对性的建模和数据增强。",
            "",
            "### 6.3 未来工作",
            "",
            "后续可以从以下方向改进。首先，采集具有像素级标注和真实 depth 的校园室内数据，用于真实场景定量验证。其次，改进深度融合方式，例如使用更稳健的深度区域编码或轻量跨模态门控，而不是仅依赖 Sobel 边界。再次，可以引入更细粒度的空间占用规则，使风险提示结合场景类型、通道方向和障碍物位置。最后，可以将系统扩展为轻量化巡检原型工具，实现批量图片管理、结果导出和人工复核功能。",
            "",
        ]
    )


def add_references(lines: list[str]) -> None:
    lines.extend(["## 参考文献", ""])
    lines.extend(REFERENCES)
    lines.append("")


def add_appendices(lines: list[str]) -> None:
    lines.extend(
        [
            "## 附录A 核心代码说明",
            "",
            "项目核心代码位于 `src/` 与 `scripts/`。`src/models/campus_depthseg_lite.py` 定义四种模型变体和前向逻辑；`src/models/mini_encoder.py` 实现轻量四阶段编码器；`src/models/depth_boundary_fusion.py` 实现 Sobel 深度边界残差融合；`src/models/weighted_fpn_decoder.py` 实现 Weighted-FPN 解码器；`src/utils/inspection.py` 实现空间占用指标和风险等级；`scripts/train.py`、`scripts/evaluate.py`、`scripts/predict_folder.py` 和 `scripts/predict_campus_demo.py` 分别对应训练、测试、测试集可视化和自采集校园定性演示。",
            "",
            "## 附录B 项目运行命令",
            "",
            "环境检查：",
            "",
            "```bash",
            "python -m compileall src scripts tests",
            "pytest -q",
            "```",
            "",
            "训练 RGBD-concat：",
            "",
            "```bash",
            "python scripts/train.py --data_dir data/nyu5 --variant rgbd_concat --experiment_name exp02_rgbd_concat_e20 --accelerator gpu --devices 1 --batch_size 4 --max_epochs 20",
            "```",
            "",
            "测试集评估：",
            "",
            "```bash",
            "python scripts/evaluate.py --data_dir data/nyu5 --split_file data/nyu5/splits/test.txt --checkpoint outputs/runs/exp02_rgbd_concat_e20/checkpoints/best.ckpt --variant rgbd_concat --batch_size 4 --accelerator gpu --devices 1",
            "```",
            "",
            "自采集校园 RGB 定性演示：",
            "",
            "```bash",
            "python scripts/predict_campus_demo.py --rgb_dir data/campus_demo/rgb --checkpoint outputs/runs/exp01_rgb_e20/checkpoints/best.ckpt --variant rgb --out_dir outputs/report_assets/campus_demo --num_samples 8",
            "```",
            "",
        ]
    )


def figure_md(number: str, title: str, path: str) -> str:
    if (PROJECT_ROOT / path).exists():
        return f"![{number} {title}]({path})\n\n{number} {title}"
    return f"[待插入：{path}]\n\n{number} {title}"


def table_md(number: str, title: str, headers: list[str], rows: list[list[str]]) -> str:
    lines = [f"{number} {title}", "", "| " + " | ".join(headers) + " |"]
    lines.append("|" + "|".join([" --- " for _ in headers]) + "|")
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def build_docx(markdown: str) -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc_placeholder(doc)
    add_doc_body(doc)
    doc.save(REPORT_DOCX)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 3"].font.size = Pt(12)


def add_cover(doc: Document) -> None:
    add_centered(doc, "湖北大学本科课程设计", 22, bold=True, space_after=48)
    add_centered(doc, "课程名称     《人工智能技术与应用》课程设计", 14, space_after=24)
    add_centered(doc, "课设题目     面向校园室内空间巡检的轻量 RGB-D 语义分割与空间占用分析系统", 14, space_after=24)
    add_centered(doc, "姓    名  易唯      学    号  202231116020106", 14, space_after=24)
    add_centered(doc, "专业年级     计算机科学与技术2022级", 14, space_after=24)
    add_centered(doc, "指导教师/职称     王雷春/副教授", 14, space_after=72)
    add_centered(doc, "2026年6月", 14, space_after=0)
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_toc_placeholder(doc: Document) -> None:
    doc.add_heading("目 录", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("请在 Word 中使用“引用-目录-更新目录”生成或更新自动目录。")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_doc_body(doc: Document) -> None:
    doc.add_heading("摘 要", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(
        doc,
        "针对校园教室、实验室、走廊等室内空间中通道遮挡、物品堆放和人工巡检主观性较强的问题，本文设计并实现了一套面向校园室内空间巡检的轻量 RGB-D 语义分割与空间占用分析系统。系统以 NYUDepthV2 数据集为基础，将原始室内场景标签映射为 other、floor、wall、obstacle、door_window 五类校园空间语义，并构建 RGB-only、RGBD-concat、RGBD-boundary 和 RGBD-concat-boundary 四种模型变体进行对比。实验结果表明，RGBD-concat 在测试集上取得最佳整体表现，mIoU 为 0.4683；相较 RGB-only 提升 4.57 个百分点。RGBD-concat-boundary 的 obstacle IoU 最高，为 0.4899，说明深度边界信息对障碍区域有一定帮助，但整体 mIoU 未超过 RGBD-concat。系统进一步实现空间占用图、障碍区域框和风险提示，并在自采集校园 RGB 图像上进行定性展示。",
    )
    add_paragraph(doc, "关键词：RGB-D 语义分割；校园巡检；空间占用分析；轻量网络；多尺度融合")
    add_heading(doc, "1 引言", 1)
    add_subsection_text(
        doc,
        "1.1 背景与意义",
        [
            "校园室内空间包含教室、实验室、宿舍楼公共区域、图书馆和走廊等多类场景。这些区域在日常使用中可能出现临时堆放物品、通道遮挡、门窗区域被遮挡、地面可见区域不足等现象。传统人工巡检通常依赖巡检人员的主观观察，记录粒度不稳定，难以长期形成结构化的空间占用分析结果。",
            "与普通图像分类不同，语义分割需要对每个像素进行类别预测，因此更适合描述空间区域的占用关系。RGB 图像包含纹理和颜色信息，depth 图像包含几何距离信息，二者结合能够更好地区分地面、墙面和家具等室内结构。",
        ],
    )
    add_subsection_text(
        doc,
        "1.2 国内外研究现状",
        [
            "语义分割任务从早期基于手工特征和传统分类器的方法，逐渐发展到以全卷积网络、编码器-解码器结构和多尺度特征融合为主的深度学习方法。FCN、U-Net、DeepLab 和 FPN 等方法分别从端到端像素预测、跳跃连接、多尺度上下文和特征金字塔等角度推动了语义分割的发展。",
            "在 RGB-D 室内语义分割方向，NYUDepthV2 是常用公开数据集之一。Depth 可以补充 RGB 难以表达的几何结构，但深度信息的使用方式需要结合模型复杂度和任务目标进行取舍。",
        ],
    )
    add_subsection_text(
        doc,
        "1.3 本文主要工作",
        [
            "本文完成五类校园空间标签映射、轻量 RGB-D 语义分割模型设计、四种融合策略消融实验、空间占用分析可视化和自采集校园场景定性展示。系统重点强调轻量、可解释和可复现，而不是依赖大模型或复杂外部工程。",
        ],
    )
    add_subsection_text(
        doc,
        "1.4 组织结构",
        [
            "本文后续章节依次介绍相关理论与技术、问题定义与数据集、模型设计与系统实现、实验结果与分析，以及总结与展望。",
        ],
    )

    add_heading(doc, "2 相关理论与技术", 1)
    add_subsection_text(doc, "2.1 语义分割任务", ["语义分割是像素级分类任务，模型需要为每个像素预测语义类别。该任务能够保留区域形状和边界信息，适合室内空间区域分析。"])
    add_subsection_text(doc, "2.2 RGB-D 多模态信息", ["RGB 提供纹理和外观信息，depth 提供空间几何信息。二者结合有助于区分地面、墙面和障碍物等室内结构。"])
    add_subsection_text(doc, "2.3 轻量多尺度特征提取", ["本文使用四阶段轻量编码器，默认通道为 [48, 96, 192, 384]，通过下采样卷积和轻量残差块提取多尺度特征。"])
    add_subsection_text(doc, "2.4 FPN 多尺度解码", ["Weighted-FPN 解码器将四个尺度特征统一到 128 通道，并使用可学习权重进行融合，兼顾局部细节和高层语义。"])
    add_subsection_text(doc, "2.5 损失函数与评价指标", ["训练损失采用 CrossEntropyLoss + DiceLoss，评价指标包括 Pixel Acc、Mean Acc、per-class IoU 和 mIoU。mIoU 是本文主要比较指标。"])

    add_heading(doc, "3 问题定义与数据集", 1)
    add_subsection_text(doc, "3.1 问题定义", ["对于 NYUDepthV2 实验，输入为 RGB 图像和 depth 图像，输出为五类语义分割 mask；对于自采集校园演示，仅使用 RGB-only checkpoint 进行定性展示。"])
    add_subsection_text(doc, "3.2 NYUDepthV2 数据集", ["本文从 nyu_depth_v2_labeled.mat 导出 1449 张 RGB-D 样本，训练集 1014 张、验证集 217 张、测试集 218 张，输入尺寸为 240 x 320。"])
    add_subsection_text(doc, "3.3 五类校园空间标签映射", ["五类标签为 other、floor、wall、obstacle、door_window，分别对应其他区域、地面、墙面、障碍物和门窗区域。"])
    add_subsection_text(doc, "3.4 数据预处理与划分", ["预处理包括 RGB 读取、depth 归一化、标签映射、resize、随机水平翻转和标准化。所有几何变换在 RGB、depth 和 label 上同步执行。"])
    add_subsection_text(doc, "3.5 类别分布分析", ["类别分布存在不均衡，door_window 面积相对较小，容易受到遮挡和光照影响。"])
    add_figures(doc, ASSET_FIGURES[:4])

    add_heading(doc, "4 模型设计与系统实现", 1)
    add_subsection_text(doc, "4.1 系统总体流程", ["系统流程包括数据准备、模型推理、语义分割、空间占用分析和可视化输出。预测 mask 被进一步用于计算占用比例、障碍区域框和风险提示。", "[待插入：系统总体流程图]"])
    add_subsection_text(doc, "4.2 CampusDepthSegLite 模型结构", ["CampusDepthSegLite 由轻量四阶段编码器、Weighted-FPN 解码器和可选深度边界残差融合模块组成。编码器通道为 [48, 96, 192, 384]，解码器通道为 128。"])
    add_subsection_text(doc, "4.3 RGB-D 融合策略", ["本文比较 RGB-only、RGBD-concat、RGBD-boundary 和 RGBD-concat-boundary 四种变体，分别验证 RGB、完整 depth、深度边界和二者组合的作用。"])
    add_subsection_text(doc, "4.4 深度边界残差融合模块", ["该模块使用 Sobel 卷积从 depth 提取边界响应，并通过 1 x 1 卷积和可学习 alpha 残差注入多尺度特征。"])
    add_subsection_text(doc, "4.5 空间占用分析模块", ["空间占用模块根据预测 mask 计算 floor_ratio、lower_floor_ratio、obstacle_ratio 和 risk_score，并使用连通域分析生成障碍区域框。风险评分是启发式指标，不能直接替代真实安全巡检决策。"])
    add_subsection_text(doc, "4.6 实验环境与超参数", ["优化器为 AdamW，学习率为 1e-4，batch size 为 4，每个实验训练 20 epoch，checkpoint 监控 val_mIoU。"])
    add_table(doc, "表4-1 模型复杂度统计", ["Variant", "Params", "Params (M)"], COMPLEXITY_ROWS)

    add_heading(doc, "5 实验结果与分析", 1)
    add_subsection_text(doc, "5.1 训练过程分析", ["RGBD-concat 作为整体最佳模型，其训练损失、验证 mIoU 和验证 Pixel Acc 曲线用于展示训练收敛趋势。"])
    add_figures(doc, ASSET_FIGURES[4:7])
    add_subsection_text(doc, "5.2 主实验结果", ["RGBD-concat 在测试集上取得最高 mIoU，为 0.4683。相比 RGB-only，mIoU 提升 4.57 个百分点，说明完整 depth 信息对室内空间分割有效。"])
    add_table(doc, "表5-1 四种模型变体测试集结果", ["Method", "Input", "Fusion", "mIoU", "Pixel Acc", "Mean Acc", "Test Loss"], EXPERIMENT_ROWS)
    add_subsection_text(doc, "5.3 分类别 IoU 分析", ["RGBD-concat 在 floor、wall 和 door_window 等类别上表现较好；RGBD-concat-boundary 的 obstacle IoU 最高，为 0.4899。door_window 类整体较难，容易与 wall 和 other 混淆。"])
    add_table(doc, "表5-2 四种模型变体 per-class IoU", ["Method", "other", "floor", "wall", "obstacle", "door_window"], PER_CLASS_ROWS)
    add_subsection_text(doc, "5.4 消融实验讨论", ["完整 depth 信息比单纯深度边界残差信息更稳定。Boundary fusion 对障碍区域有帮助，但整体没有超过 RGBD-concat，不能写成整体最佳。"])
    add_subsection_text(doc, "5.5 可视化结果分析", ["方法对比图显示 RGBD-concat 对地面和墙面连续区域预测更稳定，RGBD-concat-boundary 在部分障碍区域上响应更明显。"])
    add_figures(doc, ASSET_FIGURES[7:8])
    add_subsection_text(doc, "5.6 混淆矩阵与错误分析", ["混淆矩阵和错误案例显示，door_window 容易与 wall 和 other 混淆；遮挡严重、物体堆叠或光照不均会增加错误预测。"])
    add_figures(doc, ASSET_FIGURES[8:11])
    add_subsection_text(doc, "5.7 自采集校园场景定性展示", ["自采集图像来自图书馆、实验室、教室、走廊和遮挡较多室内场景。该部分没有像素级 GT，也没有真实 depth，不参与训练、验证或测试，不计算 mIoU、Pixel Acc 或 Mean Acc。该部分使用 RGB-only checkpoint，仅展示 Prediction mask、Occupancy map、Risk boxes、Metrics block 和 Text summary 等可视化流程。"])
    add_figures(doc, ASSET_FIGURES[11:])

    add_heading(doc, "6 总结与展望", 1)
    add_subsection_text(doc, "6.1 工作总结", ["本文完成了五类校园空间标签映射、轻量 RGB-D 分割模型设计、四种融合策略消融实验、空间占用分析和自采集校园场景定性展示。"])
    add_subsection_text(doc, "6.2 不足与局限性", ["NYUDepthV2 与真实校园场景存在 domain gap；自采集图像无 GT 和真实 depth；风险评分为启发式规则；door_window 类仍较难分割。"])
    add_subsection_text(doc, "6.3 未来工作", ["未来可采集带真实 depth 和像素级标注的校园数据，改进深度融合方式，并结合人工复核构建更完整的巡检原型系统。"])

    doc.add_heading("参考文献", level=1)
    for ref in REFERENCES:
        add_paragraph(doc, ref)

    doc.add_heading("附录A 核心代码说明", level=1)
    add_paragraph(doc, "核心代码包括模型定义、数据读取、训练评估、可视化和校园定性演示脚本。")
    doc.add_heading("附录B 项目运行命令", level=1)
    add_code_block(doc, "python -m compileall src scripts tests\npytest -q\npython scripts/predict_campus_demo.py --rgb_dir data/campus_demo/rgb --checkpoint outputs/runs/exp01_rgb_e20/checkpoints/best.ckpt --variant rgb --out_dir outputs/report_assets/campus_demo --num_samples 8")


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(8)
    else:
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)


def add_subsection_text(doc: Document, heading: str, paragraphs: Iterable[str]) -> None:
    add_heading(doc, heading, 2)
    for text in paragraphs:
        add_paragraph(doc, text)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def add_code_block(doc: Document, text: str) -> None:
    for line in text.splitlines():
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def add_centered(doc: Document, text: str, size: int, bold: bool = False, space_after: int = 12) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_figures(doc: Document, figures: Iterable[Figure]) -> None:
    for figure in figures:
        path = PROJECT_ROOT / figure.path
        if figure.page_break_before:
            doc.add_page_break()
        if path.exists():
            width_inches = figure.width_inches
            with Image.open(path) as image:
                image_width, image_height = image.size
            height_inches = width_inches * image_height / image_width
            if height_inches > figure.max_height_inches:
                height_inches = figure.max_height_inches
                width_inches = height_inches * image_width / image_height

            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_keep_with_next(paragraph)
            run = paragraph.add_run()
            run.add_picture(str(path), width=Inches(width_inches))
        else:
            add_paragraph(doc, f"[待插入：{figure.path}]")
        caption = doc.add_paragraph(f"{figure.number} {figure.title}")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(9)
    set_table_borders(table)
    doc.add_paragraph()


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


if __name__ == "__main__":
    main()
